from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT


class DocumentExportService:
    """
    Serviço para exportação de análises políticas em múltiplos formatos.
    Suporta PDF, DOCX e Markdown.
    """

    def generate_docx(self, content: str, metadata: dict) -> BytesIO:
        """
        Gera documento DOCX formatado com a análise política.
        
        Args:
            content: Conteúdo da análise em Markdown
            metadata: Metadados (político, lei, data, tempo de execução, etc)
        
        Returns:
            BytesIO com o documento DOCX gerado
        """
        document = Document()
        
        # Cabeçalho
        title = document.add_heading('📊 POSICIONÔMETRO POLÍTICO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitle = document.add_heading('Análise de Coerência Política', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph()
        
        # Metadados em tabela
        table = document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Político Analisado:'
        table.cell(0, 1).text = metadata.get('politico', 'N/A')
        
        table.cell(1, 0).text = 'Lei/Projeto:'
        table.cell(1, 1).text = metadata.get('lei', 'N/A')
        
        table.cell(2, 0).text = 'Data da Análise:'
        table.cell(2, 1).text = metadata.get('data', 'N/A')
        
        table.cell(3, 0).text = 'Tempo de Execução:'
        table.cell(3, 1).text = str(metadata.get('tempo_execucao', 'N/A'))
        
        table.cell(4, 0).text = 'Gerado em:'
        table.cell(4, 1).text = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        
        document.add_paragraph()
        document.add_paragraph('_' * 50)
        document.add_paragraph()
        
        # Conteúdo da análise
        heading = document.add_heading('Resultado da Análise:', level=1)
        
        # Processar markdown para texto formatado
        lines = content.split('\n')
        for line in lines:
            if not line.strip():
                continue
                
            # Títulos
            if line.startswith('### '):
                document.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '):
                document.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                document.add_heading(line.replace('# ', ''), level=1)
            # Listas
            elif line.startswith('- '):
                p = document.add_paragraph(line.replace('- ', ''), style='List Bullet')
            elif line.strip().startswith('* '):
                p = document.add_paragraph(line.strip().replace('* ', ''), style='List Bullet')
            # Separadores
            elif line.strip() == '---':
                document.add_paragraph('_' * 50)
            # Texto normal
            else:
                # Processar negrito simples
                text = line
                if '**' in text:
                    p = document.add_paragraph()
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            p.add_run(part)
                        else:
                            p.add_run(part).bold = True
                else:
                    document.add_paragraph(text)
        
        # Rodapé
        document.add_paragraph()
        document.add_paragraph('_' * 50)
        footer = document.add_paragraph()
        footer.add_run('Documento gerado automaticamente pelo Posicionômetro Político').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em BytesIO
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_pdf(self, content: str, metadata: dict) -> BytesIO:
        """
        Gera documento PDF formatado com a análise política.
        
        Args:
            content: Conteúdo da análise em Markdown
            metadata: Metadados (político, lei, data, tempo de execução, etc)
        
        Returns:
            BytesIO com o documento PDF gerado
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#003366'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=HexColor('#333333'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#003366'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        # Elementos do PDF
        story = []
        
        # Título
        story.append(Paragraph("📊 POSICIONÔMETRO POLÍTICO", title_style))
        story.append(Paragraph("Análise de Coerência Política", subtitle_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Metadados
        story.append(Paragraph("<b>Político Analisado:</b> " + metadata.get('politico', 'N/A'), normal_style))
        story.append(Paragraph("<b>Lei/Projeto:</b> " + metadata.get('lei', 'N/A'), normal_style))
        story.append(Paragraph("<b>Data da Análise:</b> " + metadata.get('data', 'N/A'), normal_style))
        story.append(Paragraph("<b>Tempo de Execução:</b> " + str(metadata.get('tempo_execucao', 'N/A')), normal_style))
        story.append(Paragraph("<b>Gerado em:</b> " + datetime.now().strftime('%d/%m/%Y %H:%M:%S'), normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Linha divisória
        story.append(Paragraph("_" * 80, normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Conteúdo
        story.append(Paragraph("Resultado da Análise:", heading_style))
        story.append(Spacer(1, 0.1*inch))
        
        # Processar markdown para PDF
        lines = content.split('\n')
        for line in lines:
            if not line.strip():
                story.append(Spacer(1, 0.1*inch))
                continue
            
            # Títulos
            if line.startswith('### '):
                story.append(Paragraph(line.replace('### ', ''), heading_style))
            elif line.startswith('## '):
                story.append(Paragraph(line.replace('## ', ''), heading_style))
            elif line.startswith('# '):
                story.append(Paragraph(line.replace('# ', ''), title_style))
            # Separadores
            elif line.strip() == '---':
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph("_" * 80, normal_style))
                story.append(Spacer(1, 0.1*inch))
            # Texto normal
            else:
                # Converter negrito markdown para HTML
                text = line.replace('**', '<b>', 1).replace('**', '</b>', 1) if '**' in line else line
                story.append(Paragraph(text, normal_style))
        
        # Rodapé
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("_" * 80, normal_style))
        footer_style = ParagraphStyle(
            'Footer',
            parent=normal_style,
            fontSize=9,
            textColor=HexColor('#808080'),
            alignment=TA_CENTER
        )
        story.append(Paragraph("<i>Documento gerado automaticamente pelo Posicionômetro Político</i>", footer_style))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_markdown(self, content: str, metadata: dict) -> BytesIO:
        """
        Gera arquivo Markdown formatado com a análise política.
        
        Args:
            content: Conteúdo da análise em Markdown
            metadata: Metadados (político, lei, data, tempo de execução, etc)
        
        Returns:
            BytesIO com o arquivo Markdown gerado
        """
        md_content = []
        
        # Cabeçalho
        md_content.append("# 📊 POSICIONÔMETRO POLÍTICO\n")
        md_content.append("## Análise de Coerência Política\n")
        md_content.append("\n---\n\n")
        
        # Metadados
        md_content.append("### Informações da Análise\n\n")
        md_content.append(f"- **Político Analisado:** {metadata.get('politico', 'N/A')}\n")
        md_content.append(f"- **Lei/Projeto:** {metadata.get('lei', 'N/A')}\n")
        md_content.append(f"- **Data da Análise:** {metadata.get('data', 'N/A')}\n")
        md_content.append(f"- **Tempo de Execução:** {metadata.get('tempo_execucao', 'N/A')}\n")
        md_content.append(f"- **Gerado em:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        md_content.append("\n---\n\n")
        
        # Conteúdo
        md_content.append("## Resultado da Análise\n\n")
        md_content.append(content)
        md_content.append("\n\n---\n\n")
        
        # Rodapé
        md_content.append("_Documento gerado automaticamente pelo Posicionômetro Político_\n")
        
        # Criar buffer
        buffer = BytesIO()
        buffer.write('\n'.join(md_content).encode('utf-8'))
        buffer.seek(0)
        return buffer
